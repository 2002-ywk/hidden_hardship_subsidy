from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]


def find_markdown() -> Path:
    matches = sorted(ROOT.glob("*更新版.md"))
    if not matches:
        raise FileNotFoundError("未找到更新版 markdown 文档")
    return matches[0]


def build_doc(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run(stripped[2:].strip())
            r.bold = True
            r.font.name = "黑体"
            r.font.size = Pt(16)
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            r.bold = True
            r.font.name = "黑体"
            r.font.size = Pt(14)
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:].strip())
            r.bold = True
            r.font.name = "黑体"
            r.font.size = Pt(12)
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:].strip())
            continue

        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == ".":
            p = doc.add_paragraph(style="List Number")
            p.add_run(stripped[stripped.find(".") + 1 :].strip())
            continue

        doc.add_paragraph(stripped)

    doc.save(output_path)


def main() -> None:
    markdown_path = find_markdown()
    output_path = ROOT / "学生食堂消费补助需求文档-更新版.docx"
    build_doc(markdown_path.read_text(encoding="utf-8"), output_path)
    print(output_path)


if __name__ == "__main__":
    main()
